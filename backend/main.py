from fastapi import FastAPI, File, Form, UploadFile, Request, HTTPException
import chromadb
from chromadb.config import Settings
from chromadb import PersistentClient
from typing import List, Annotated, Union
from contextlib import asynccontextmanager
import pymupdf
import docx
import io
import os
import time
import asyncio
from store_to_db import *
from rank_candidates import *


app = FastAPI()

# setting up ollama so that the container can reach it
OLLAMA_HOST = os.getenv("OLLAMA_HOST", "http://host.docker.internal:11434")

file_size_limit = 2_000_000 # 2 MB
timeout_limit = 10 # Maximum given time to get the file to ollama to extract its metadata
# and move the document to the chromadb


# Create an Persistent client
# chroma_client = chromadb.PersistentClient(
#     path='../database',
#     settings=Settings(allow_reset=True)
# )

# Create an ephemeral client
chroma_client = chromadb.Client(
    settings=Settings(allow_reset=True)
)

# Create a collection if not exists to add and retrieve embeddings from
student_collection = chroma_client.create_collection(
    name="students_collection",
    embedding_function=None
)

async def retry_async_function(func, *args, timer=1, attempts=3, **kwargs):
    for attempt in range(attempts):
        print('loop')
        try:
            return await asyncio.wait_for(func(*args, **kwargs), timeout=timer)
        except:
            print('failed')
            print(f'attempt: {attempt + 1} with timer {timer}')
            print("*"*10)
            timer *= 2
    raise asyncio.TimeoutError()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    pdf_file = pymupdf.open(stream=file_bytes, filetype="pdf")
    return " ".join(page.get_text() for page in pdf_file)

def extract_text_from_docx(file_bytes: bytes) -> str:    
    docx_files = docx.Document(io.BytesIO(file_bytes))
    return ' '.join([paragraph.text for paragraph in docx_files])

def extract_text_from_single_pdf(file_bytes: bytes) -> str:
    pdf_files = pymupdf.open(stream=file_bytes, filetype="pdf")
    return " ".join(page.get_text() for page in pdf_files)

def extract_text_from_single_docx(file_bytes: bytes) -> str:
    docx_file = docx.Document(io.BytesIO(file_bytes))
    return ' '.join([paragraph.text for paragraph in docx_file.paragraphs])

@app.delete('/remove_collection_data')
async def erase_collection_data():

    try:
        chroma_client.delete_collection('students_collection')
        chroma_client.reset()
        return{
            'message': 'collection data erased successfully'
        }
    except Exception as e:
        return {
            'message': 'something went wrong!',
            'error': str(e)
        }

@app.post('/upload_cvs')
async def upload_files(files: list[UploadFile] = File(...)):

    results = []

    for file in files:

        try: 
            filename = file.filename.lower()
            content = await file.read()
            file_size = len(content)

            if filename.endswith('.pdf'):

                if file_size <= file_size_limit:
                    print("pass pdf to the next stage")
                    cv_str = extract_text_from_pdf(content)
                else:
                    print("file is too large")
                    results.append({
                        'filename': filename,
                        'status': 'failed',
                        'reason': 'File is too large'
                    })
                    continue
            elif filename.endswith('.docx'):

                if file_size <= file_size_limit:
                    print("pass word to the next stage")
                    cv_str = extract_text_from_docx(content)
                else:
                    print("file is too large")
                    results.append({
                        'filename': filename,
                        'status': 'failed',
                        'reason': 'File is too large'
                    })
                    continue
            else:
                print('unsupported file')
                results.append({
                        'filename': filename,
                        'status': 'failed',
                        'reason': 'Unsupported file type'
                    })
                continue

            # save the file in the vector database
            try:
                collection = chroma_client.get_collection('students_collection')
                try:
                    # result = await asyncio.wait_for(store_cv_in_db(collection, cv_str), timeout=timeout_limit)
                    # await asyncio.wait_for(store_cv_in_db(collection, cv_str), timeout=timeout_limit)
                    await retry_async_function(lambda: store_cv_in_db(collection, cv_str), timer=2, attempts=3)
                except asyncio.TimeoutError as e:
                    results.append({
                        'filename': filename,
                        'status': 'failed',
                        'reason': str(e)
                    })
                    continue

                results.append({
                'filename': filename,
                'status': 'success'
                })

            except Exception as e:
                results.append({
                'filename': filename,
                'status': 'failed',
                'reason': str(e)
            })


        except Exception as e:

            print(f"error: \n{str(e)}\n")

            results.append({
                'filename': filename,
                'status': 'failed',
                'reason': str(e)
            })


    total = len(results)
    success = sum(1 for r in results if r['status'] == 'success')
    failed = total - success

    return {
        'total': total,
        'success': success,
        'failed': failed,
        'details': results
    }

@app.get('/collection_count')
async def get_collection_count():

    try:
        collection = chroma_client.get_collection(name='students_collection')
        record_count = collection.count()
        if record_count > 0:
            return {
                'status': 'success',
                'collection_count' : record_count
            }
        else:
            return {
                'status': 'failed',
                'collection_count': 'Collection is empty'
            }
    except Exception as e:
        return {
            'status': 'failed',
            'collection_count': str(e)
        }
    
# @app.get('/list_collections')
# async def list_collections():
    try:
        list_collections = chroma_client.list_collections()
        return {
            'status': 'success',
            'collections': list_collections
        }
    except Exception as e:
        return {
            'status': 'failed',
            'reason': str(e)
        }


@app.get('/collection_peek')
async def peek():
    try:
        collection = chroma_client.get_collection('students_collection')
        data = collection.peek(3)
        return{
            'ids': data['ids'],
            'documents': data['documents'],
            'metadatas': data['metadatas']
        }
    except Exception as e:
        return{
            'error': str(e)
        }

@app.post('/find_candidates')
async def find_candidates(job_desc: UploadFile = File(...), n_top_applicants: int = Form(...)):
        
        try:
            filename = job_desc.filename.lower()
            content = await job_desc.read()
            file_size = len(content)

            if filename.endswith('.pdf'):
                if file_size < file_size_limit:
                    job_desc_str = extract_text_from_single_pdf(content)
                    try:
                        collection = chroma_client.get_collection('students_collection')
                        top_applicants_result = await retry_async_function(lambda: get_top_applicants(collection, job_desc_str, n_top_applicants))
                        # top_applicants_result = get_top_applicants(collection, job_desc_str, n_top_applicants)
                        return{
                            'status': 'success',
                            'result': top_applicants_result
                        }
                    
                    except Exception as e:
                        return{
                            'status': 'failed',
                            'result': "could not get top applicants from pdf"
                        }

                else:
                    print("file is too large")
                    return{
                        'filename': filename,
                        'status': 'failed',
                        'reason': 'File is too large'
                    }
                                
            elif filename.endswith('.docx'):
                if file_size < file_size_limit:
                    job_desc_str = extract_text_from_single_docx(content)
                    try:
                        collection = chroma_client.get_collection('students_collection')
                        top_applicants_result = await retry_async_function(lambda: get_top_applicants(collection, job_desc_str, n_top_applicants))
                        # top_applicants = get_top_applicants(collection, job_desc_str, n_top_applicants)
                        return{
                            'status': 'success',
                            'result': top_applicants_result
                        }
                    except Exception as e:
                        return{
                            'status': 'failed',
                            'result': "could not get top applicants from word"
                        }
                else:
                    print("file is too large")
                    return{
                        'filename': filename,
                        'status': 'failed',
                        'reason': 'File is too large'
                    }
            else:
                return{
                    'filename': filename,
                    'status': 'failed',
                    'reason': 'Unsupported file type'
                }

        except Exception as e:
            return{
                'status': 'failed',
                'reason': str(e)
            }
        




